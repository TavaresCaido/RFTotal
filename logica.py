import os
from docx import Document
from docx.shared import Cm
from PIL import Image
import openai
from ultralytics import YOLO
import logging

openai.api_key = os.getenv("OPENAI_API_KEY")

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class Logica:
    def __init__(self):
        self.diretorio_imagens = ""
        self.arquivo_word = ""
        self.modelo_yolo = YOLO("yolov8n.pt")

    def selecionar_diretorio(self):
        self.diretorio_imagens = filedialog.askdirectory()
        return self.diretorio_imagens

    def selecionar_arquivo_word(self):
        self.arquivo_word = filedialog.askopenfilename(filetypes=[("Documentos Word", "*.docx")])
        return self.arquivo_word

    def extrair_texto_word(self):
        try:
            doc = Document(self.arquivo_word)
            textos = [p.text for p in doc.paragraphs if p.text.strip()]
            return textos
        except Exception as e:
            logging.error(f"Erro ao extrair texto do Word: {e}")
            return []

    def identificar_imagens_por_texto(self, textos):
        try:
            resposta = openai.ChatCompletion.create(
                model="gpt-4-turbo",
                messages=[
                    {"role": "system", "content": "Você é um engenheiro que analisa relatórios de obras e identifica as imagens necessárias."},
                    {"role": "user", "content": f"Baseado nesses comentários de obra, quais tipos de imagens são necessárias?\n{textos}"}
                ]
            )
            return resposta['choices'][0]['message']['content'].split("\n")
        except Exception as e:
            logging.error(f"Erro ao identificar imagens por texto: {e}")
            return []

    def encontrar_imagens_corretas(self, descricao):
        imagens_correspondentes = []
        for img in os.listdir(self.diretorio_imagens):
            caminho = os.path.join(self.diretorio_imagens, img)
            if not caminho.lower().endswith((".png", ".jpg", ".jpeg")):
                continue
            try:
                resultados = self.modelo_yolo(caminho)
                for r in resultados:
                    objetos_detectados = [self.modelo_yolo.names[int(c)] for c in r.boxes.cls]
                    if descricao.lower() in objetos_detectados:
                        imagens_correspondentes.append(caminho)
            except Exception as e:
                logging.error(f"Erro ao processar imagem {caminho}: {e}")
        return imagens_correspondentes

    def redimensionar_imagem(self, imagem_path, altura_desejada=8.8):
        try:
            img = Image.open(imagem_path)
            proporcao = (altura_desejada * 37.795) / img.height
            largura_nova = int(img.width * proporcao)
            altura_nova = int(altura_desejada * 37.795)
            img = img.resize((largura_nova, altura_nova))
            novo_caminho = f"temp_{os.path.basename(imagem_path)}"
            img.save(novo_caminho)
            return novo_caminho
        except Exception as e:
            logging.error(f"Erro ao redimensionar imagem {imagem_path}: {e}")
            return None

    def inserir_imagens_no_word(self):
        if not self.diretorio_imagens or not self.arquivo_word:
            return "Erro: Selecione um diretório e um arquivo Word antes de continuar!"

        try:
            doc = Document(self.arquivo_word)
            textos = self.extrair_texto_word()
            descricoes_imagens = self.identificar_imagens_por_texto(textos)

            imagens_para_inserir = {}
            for descricao in descricoes_imagens:
                imagens_para_inserir[descricao] = self.encontrar_imagens_corretas(descricao)

            for paragrafo in doc.paragraphs:
                for descricao, imagens in imagens_para_inserir.items():
                    if descricao.lower() in paragrafo.text.lower() and imagens:
                        imagem_corrigida = self.redimensionar_imagem(imagens[0])
                        if imagem_corrigida:
                            paragrafo.add_run().add_picture(imagem_corrigida, height=Cm(8.8))

            novo_arquivo = "relatorio_final.docx"
            doc.save(novo_arquivo)
            return f"✅ Relatório gerado com sucesso: {novo_arquivo}"
        except Exception as e:
            logging.error(f"Erro ao inserir imagens no Word: {e}")
            return f"Ocorreu um erro ao gerar o relatório: {e}"